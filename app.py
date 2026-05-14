import streamlit as st
import os
from io import BytesIO
from docx import Document
from utils.document_reader import (
    extract_text_from_pdf,
    extract_text_from_ppt,
    extract_text_from_image
)
from utils.ai_engine import ask_ai
from frontend.components import load_css, show_header, show_stats_strip, show_all_features, show_upload_prompt, show_quiz_card, show_result_card, show_score_card

# Page config
st.set_page_config(
    page_title="AI Study Companion",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Load CSS
load_css()

# Show header
show_header()
show_stats_strip()

# ================= SAVE FUNCTION =================
# ================= SIMPLE SAVE FUNCTION =================
def save_document(content, title):
    """Simple save function - direct download without any options"""
    
    # Create filename
    filename = title.lower().replace(" ", "_").replace("___", "_")
    
    # Create Word document
    doc = Document()
    doc.add_heading(title.replace("_", " "), level=1)
    
    # Add content with proper formatting
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()  # Empty line for spacing
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Direct download button
    st.download_button(
        label="📥 Download as Word Document",
        data=buffer,
        file_name=f"{filename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

# ================= RESET OUTPUT =================
if 'last_option' not in st.session_state:
    st.session_state.last_option = None

def reset_output(option):
    if st.session_state.last_option != option:
        st.session_state.output = ""
        st.session_state.quiz_answers = {}
    st.session_state.last_option = option

# ================= FILE UPLOADER =================
uploaded_file = st.file_uploader(
    "Choose a file",
    type=["pdf", "pptx", "png", "jpg", "jpeg"],
    label_visibility="collapsed",
    help="Upload PDF, PPTX, PNG, JPG, JPEG (Max 200MB)"
)

# Show upload prompt if no file
if not uploaded_file:
    # Clean upload prompt
    st.markdown("""
        <div style="text-align: center; padding: 3rem 2rem; background: linear-gradient(135deg, #0a0a12, #08080e); border-radius: 1rem; border: 1px solid rgba(124, 58, 237, 0.3); margin: 1rem 0;">
            <div style="font-size: 3rem; margin-bottom: 1rem;">📚</div>
            <h3 style="color: #e8e4ff; margin-bottom: 0.5rem;">Ready to Study?</h3>
            <p style="color: #a78bfa; margin-bottom: 0;">Click the button to upload your document</p>
        </div>
    """, unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("### 🚀 What you can do with AI Study Companion")
    show_all_features()
    st.stop()

# Process uploaded file
file_type = uploaded_file.type

with st.spinner("📄 Processing your document..."):
    if "pdf" in file_type:
        text = extract_text_from_pdf(uploaded_file)
    elif "presentation" in file_type:
        text = extract_text_from_ppt(uploaded_file)
    elif "image" in file_type:
        text = extract_text_from_image(uploaded_file)
    else:
        st.error("Unsupported file type")
        st.stop()

if len(text.strip()) == 0:
    st.error("No readable text found in the document")
    st.stop()

st.success(f"✅ Document processed successfully! ({len(text)} characters extracted)")

# ================= DROPDOWN SELECT BOX =================
st.markdown('<label style="color: #a78bfa; font-weight: 500; margin-bottom: 10px; display: block;">📌 What would you like to do?</label>', unsafe_allow_html=True)

# Simple 2 columns radio without labels
option = st.radio(
    "Select option",
    ["📝 Summary", "📖 Detailed Overview", "❓ Generate MCQs", "📋 Short Questions", "⭐ Key Points", "📚 Difficult Words"],
    label_visibility="collapsed",
    horizontal=False
)

# Clean option name
if option.startswith("📝"):
    clean_option = "Summary"
elif option.startswith("📖"):
    clean_option = "Detailed"
elif option.startswith("❓"):
    clean_option = "MCQs"
elif option.startswith("📋"):
    clean_option = "Short"
elif option.startswith("⭐"):
    clean_option = "Key"
elif option.startswith("📚"):
    clean_option = "Difficult"
else:
    clean_option = option

# ================= CONTENT DISPLAY HELPER =================
def display_content(content, title):
    show_result_card(title)
    st.markdown(f'<div style="background: #0e0e1a; padding: 1.2rem; border-radius: 0.75rem; border: 1px solid rgba(139, 92, 246, 0.3); color: #d1d5db; line-height: 1.6;">', unsafe_allow_html=True)
    st.markdown(content)
    st.markdown('</div>', unsafe_allow_html=True)
    save_document(content, title.replace(" ", "_"))

# ================= SUMMARY =================
if clean_option == "Summary":
    if st.button("✨ Generate Summary", use_container_width=True):
        with st.spinner("🤖 AI is analyzing and summarizing your document..."):
            prompt = f"Summarize this text in bullet points (max 500 words):\n\n{text[:4000]}"
            response = ask_ai(prompt)
        st.session_state.output = response
    
    if st.session_state.get("output"):
        display_content(st.session_state.output, "Summary")

# ================= DETAILED OVERVIEW =================
elif clean_option == "Detailed":
    if st.button("✨ Generate Detailed Overview", use_container_width=True):
        with st.spinner("🤖 AI is creating detailed explanation..."):
            prompt = f"Provide a detailed explanation of this topic with introduction, key concepts, examples, and conclusion:\n\n{text[:4000]}"
            response = ask_ai(prompt)
        st.session_state.output = response
    
    if st.session_state.get("output"):
        display_content(st.session_state.output, "Detailed_Overview")

# ================= MCQs =================
elif clean_option == "MCQs":
    num = st.slider("Number of MCQs", 1, 10, 5)
    
    if st.button("✨ Generate MCQs", use_container_width=True):
        with st.spinner("🤖 AI is generating questions..."):
            prompt = f"Generate {num} MCQs from this text. Format each as:\nQ1. Question\nA) option\nB) option\nC) option\nD) option\nAnswer: X\n\nText:\n{text[:4000]}"
            response = ask_ai(prompt)
        
        # Parse MCQs
        mcqs = {}
        lines = response.splitlines()
        q, opts, ans = "", {}, ""
        
        for line in lines:
            line = line.strip()
            if line.startswith("Q") and len(line) > 1 and line[1].isdigit():
                if q:
                    mcqs[q] = {"options": opts, "answer": ans}
                q = line
                opts, ans = {}, ""
            elif line.startswith(("A)", "B)", "C)", "D)")):
                if len(line) > 2:
                    opts[line[0]] = line[3:].strip()
            elif line.startswith("Answer:"):
                ans = line.split(":")[1].strip()
        if q:
            mcqs[q] = {"options": opts, "answer": ans}
        
        st.session_state.quiz_answers = mcqs
    
    if st.session_state.get("quiz_answers") and st.session_state.quiz_answers:
        st.markdown("---")
        st.subheader("📝 Take the Quiz")
        
        user_answers = {}
        for idx, (question, data) in enumerate(st.session_state.quiz_answers.items()):
            if data['options']:
                user_answers[question] = show_quiz_card(question, data['options'], idx)
        
        if user_answers and st.button("✅ Submit Answers", use_container_width=True):
            correct = sum(
                1 for q, ans in user_answers.items()
                if q in st.session_state.quiz_answers and ans == st.session_state.quiz_answers[q]["answer"]
            )
            show_score_card(correct, len(user_answers))

# ================= SHORT QUESTIONS =================
elif clean_option == "Short":
    num = st.slider("Number of questions", 1, 15, 5)
    
    if st.button("✨ Generate Short Questions", use_container_width=True):
        with st.spinner("🤖 AI is generating questions with answers..."):
            prompt = f"Generate {num} short answer questions with answers (each answer 2-3 lines) from this text. Format as:\nQ1. Question\nAnswer: ...\n\nText:\n{text[:4000]}"
            response = ask_ai(prompt)
        st.session_state.output = response
    
    if st.session_state.get("output"):
        display_content(st.session_state.output, "Short_Questions")

# ================= KEY POINTS =================
elif clean_option == "Key":
    if st.button("✨ Extract Key Points", use_container_width=True):
        with st.spinner("🤖 AI is extracting important points..."):
            prompt = f"Extract the 10 most important key points from this text in bullet points:\n\n{text[:4000]}"
            response = ask_ai(prompt)
        st.session_state.output = response
    
    if st.session_state.get("output"):
        display_content(st.session_state.output, "Key_Points")

# ================= DIFFICULT WORDS =================
elif clean_option == "Difficult":
    if st.button("✨ Find Difficult Words", use_container_width=True):
        with st.spinner("🤖 AI is identifying difficult words..."):
            prompt = f"Identify difficult or technical words from this text and explain them in simple terms. Format as a list:\n\n{text[:4000]}"
            response = ask_ai(prompt)
        st.session_state.output = response
    
    if st.session_state.get("output"):
        display_content(st.session_state.output, "Difficult_Words")